from docx import Document
import os
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


def format_filename(title: str, max_lenght: int = 12) -> str:
    return title.replace(" ", "")[:max_lenght]


def trim_text(text, max_lenght=1500):
    return text[:max_lenght] + '...' if len(text) > max_lenght else text


def save_blog_to_docx(title, content):
    doc = Document()
    doc.add_paragraph(content)

    formatted_title = format_filename(title)

    directory = "../data/blogDocs"
    os.makedirs(directory, exist_ok=True)

    filename = os.path.join(directory, f"{formatted_title}.docx")
    doc.save(filename)
    logging.info(f"Блог {filename} сохранён в {directory}")


def read_blog_from_docx(filepath):
    doc = Document(filepath)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


def save_slogan_to_docx(title, slogan):
    doc = Document()
    doc.add_paragraph(slogan)

    formatted_title = format_filename(title)

    directory = "../data/sloganDocs"
    os.makedirs(directory, exist_ok=True)
    filename = os.path.join(directory, f"{formatted_title}_slogan.docx")
    doc.save(filename)
    logging.info(f"Слоган {filename} сохранён в {directory}")


def read_slogan_from_docx(filepath):
    doc = Document(filepath)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)
